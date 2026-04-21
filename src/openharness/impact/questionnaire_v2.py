"""DD Questionnaire v2 — conditional branching (Phase 15.6).

Extends :class:`openharness.impact.dd_checklist.DDQuestion` with
*branching rules*: a question can expose optional follow-ups that only
fire when the founder's answer satisfies a condition.

Rules engine
------------
A :class:`BranchRule` evaluates a plain-text answer and, if the
condition matches, surfaces the follow-up question(s).  The three
condition kinds cover ~90% of what's needed in impact DD:

* ``contains`` — keyword found (case-insensitive).
* ``equals`` — exact match (case-insensitive).
* ``missing`` — answer is blank / "N/A" / "don't know".

Example
-------
::

    rules = [
        BranchRule(parent_id="who-7",
                   kind="missing",
                   follow_up_ids=["who-7a", "who-7b"],
                   rationale="If Who data is missing, always ask about survey methodology."),
    ]
    active = expand_active(answered={"who-7": ""}, rules=rules)

The companion :func:`render_branching_docx` writes a Word document that
*only* includes the follow-ups that fired — which is what a GP analyst
would paste back to the founder.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, Literal

from pydantic import BaseModel, Field

from openharness.impact.dd_checklist import DDQuestion


BranchKind = Literal["contains", "equals", "missing"]


class BranchRule(BaseModel):
    parent_id: str
    kind: BranchKind
    value: str = ""
    follow_up_ids: list[str] = Field(default_factory=list)
    rationale: str = ""


def _evaluate(rule: BranchRule, answer: str) -> bool:
    ans = (answer or "").strip().lower()
    if rule.kind == "missing":
        return ans in {"", "n/a", "na", "unknown", "don't know", "dont know", "tbd"}
    val = rule.value.strip().lower()
    if rule.kind == "equals":
        return ans == val
    if rule.kind == "contains":
        return val in ans
    return False


def active_follow_up_ids(
    answers: dict[str, str],
    rules: Iterable[BranchRule],
) -> list[str]:
    """Return the ordered set of follow-up IDs that fire given ``answers``."""
    seen: set[str] = set()
    out: list[str] = []
    for rule in rules:
        ans = answers.get(rule.parent_id, "")
        if _evaluate(rule, ans):
            for fu in rule.follow_up_ids:
                if fu not in seen:
                    out.append(fu)
                    seen.add(fu)
    return out


def expand_active(
    *,
    answered: dict[str, str],
    rules: Iterable[BranchRule],
    catalogue: Iterable[DDQuestion],
) -> list[DDQuestion]:
    """Return only the follow-up questions that the branch rules activate."""
    active_ids = active_follow_up_ids(answered, rules)
    id_to_q = {q.id: q for q in catalogue}
    return [id_to_q[i] for i in active_ids if i in id_to_q]


def render_branching_docx(
    path: str | Path,
    *,
    parent_questions: Iterable[DDQuestion],
    answered: dict[str, str],
    rules: Iterable[BranchRule],
    catalogue: Iterable[DDQuestion],
    company_name: str = "Company",
) -> Path:
    """Write a ``.docx`` containing only parents + their fired follow-ups.

    Requires the optional dependency ``python-docx``.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover — optional dep
        raise ImportError(
            "Branching Word export requires 'python-docx'. "
            "Install with: pip install python-docx"
        ) from exc

    rules_list = list(rules)
    catalogue_list = list(catalogue)
    id_to_q = {q.id: q for q in catalogue_list}
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(f"DD Questionnaire v2 (branching) — {company_name}", level=0)
    doc.add_paragraph(
        "This document contains every parent question plus the branching "
        "follow-ups that were activated by the founder's initial answers. "
        "Please respond in the 'Response' field under each question."
    )

    for parent in parent_questions:
        doc.add_heading(parent.question, level=2)
        meta = doc.add_paragraph()
        meta.add_run(
            f"Category: {parent.category} · Priority: {parent.priority}"
        )
        doc.add_paragraph("Response: ______________________________________________")

        # Find follow-ups for this parent
        for rule in rules_list:
            if rule.parent_id != parent.id:
                continue
            if not _evaluate(rule, answered.get(parent.id, "")):
                continue
            for fu_id in rule.follow_up_ids:
                fu = id_to_q.get(fu_id)
                if fu is None:
                    continue
                sub = doc.add_paragraph()
                r = sub.add_run(f"    ↳ Follow-up: {fu.question}")
                r.italic = True
                if rule.rationale:
                    rr = doc.add_paragraph()
                    rr.add_run(f"    Why: {rule.rationale}").italic = True
                doc.add_paragraph("    Response: ______________________________________________")

    doc.save(str(p))
    return p


__all__ = [
    "BranchKind",
    "BranchRule",
    "active_follow_up_ids",
    "expand_active",
    "render_branching_docx",
]
