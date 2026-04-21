"""HTML + Word renderer for the **DD Questionnaire Helper**.

Takes a ``DDChecklistResult`` from
``openharness.impact.dd_checklist.analyze_document_coverage`` and produces a
self-contained, print-ready HTML page (and optionally a ``.docx`` file) that
focuses on the *work the GP analyst still needs to do*, not just the coverage
statistics:

  * **Key risk areas** — categories with the weakest coverage, sorted by
    severity (high-priority gaps first, then by coverage percentage).
  * **Priority questionnaire** — the full follow-up questionnaire to send to
    the founder/sponsor, sorted by severity *and* by the natural DD sequence
    (Thesis → Theory of Change → What → Who → How-much → Contribution → Risk
    → Measurement → Governance → Sector → Exit). Each question has space for
    an answer and flags the supporting evidence / policy documents that
    should be attached.
  * **Evidence & document gaps** — which classes of supporting documents
    (policies, KPI dashboards, audit letters, stakeholder surveys, …) the GP
    should request in a single consolidated information request.
  * **Already-addressed questions** — kept as a collapsible appendix so the
    analyst can audit the auto-match.

The shared v2 chrome (hero + KPI strip + sticky TOC + callouts) is reused
for visual consistency with the full Impact Report and the IC memo.

The Word export (:func:`render_dd_questionnaire_docx`) mirrors the HTML
questionnaire section as an editable ``.docx`` so the analyst can hand the
file directly to the founder / deal team.

The HTML renderers keep their old names (``render_dd_report_html`` /
``save_dd_report_html``) as aliases to avoid breaking downstream callers.
"""
from __future__ import annotations

from collections import Counter, defaultdict
from datetime import date
from html import escape
from pathlib import Path
from typing import Iterable

from openharness.impact.dd_checklist import (
    DDChecklistResult,
    DDQuestion,
    DDQuestionMatch,
    EVIDENCE_LEVELS,
)
from openharness.impact.report_templates.report_v2 import (
    render_footer,
    render_hero,
    render_kpi_strip,
    render_toc,
    wrap_document,
)

# ---------------------------------------------------------------------------
# Ordering / taxonomy helpers
# ---------------------------------------------------------------------------

#: Natural DD sequence — categories earlier in this list are investigated
#: first in a typical impact DD, so we order the questionnaire accordingly
#: inside each priority band.
_DD_SEQUENCE: tuple[str, ...] = (
    "impact_thesis",
    "theory_of_change",
    "what_outcomes",
    "who_stakeholders",
    "how_much_scale",
    "contribution",
    "risk",
    "negative_impact",
    "measurement_systems",
    "governance_esg",
    "sdg_alignment",
    "stakeholder_voice",
    "exit_sustainability",
    "financial_sustainability",
    "team_capability",
    "market_context",
    "product_design",
    "supply_chain",
    "investor_alignment",
)

#: Priority ranking — lower number = more severe / asked first.
_PRIORITY_RANK: dict[str, int] = {"high": 0, "medium": 1, "low": 2}

#: Keyword → evidence-request mapping. If a missing question's text contains
#: any of the keys we surface the mapped evidence ask in the information
#: request. Kept intentionally small — the analyst can always add more.
_EVIDENCE_HINTS: tuple[tuple[tuple[str, ...], str], ...] = (
    (("theory of change", "logic model", "causal chain", "pathway"),
     "Theory-of-change diagram or written logic model"),
    (("policy", "code of conduct", "guideline"),
     "Formal policy document (e.g. ESG, safeguarding, data-privacy)"),
    (("audit", "verified", "third-party"),
     "Third-party audit / verification letter"),
    (("survey", "feedback", "net promoter", "NPS", "satisfaction"),
     "Stakeholder survey instrument + anonymised results"),
    (("KPI", "dashboard", "metric", "indicator", "IRIS+"),
     "Live KPI / metric dashboard export (last 12 months)"),
    (("financial", "unit economics", "ebitda", "revenue"),
     "Audited or management financial statements"),
    (("exit", "sustainability", "long-term"),
     "Exit / sustainability plan memo"),
    (("board", "governance", "committee", "oversight"),
     "Board / impact-committee charter and minutes"),
    (("benchmark", "comparator", "peer"),
     "Peer / sector benchmark analysis"),
    (("certificat", "accredit"),
     "Certification / accreditation certificate"),
)

#: Category → human-readable label for section headings and the Word export.
_CATEGORY_LABEL: dict[str, str] = {
    "impact_thesis":          "Impact thesis & mission",
    "theory_of_change":       "Theory of change",
    "what_outcomes":          "What — intended outcomes",
    "who_stakeholders":       "Who — stakeholders & beneficiaries",
    "how_much_scale":         "How much — scale, depth, duration",
    "contribution":           "Contribution / additionality",
    "risk":                   "Impact risk",
    "negative_impact":        "Negative impact / do-no-harm",
    "measurement_systems":    "Measurement & management systems",
    "governance_esg":         "Governance & ESG",
    "sdg_alignment":          "SDG alignment",
    "stakeholder_voice":      "Stakeholder voice",
    "exit_sustainability":    "Exit & long-run sustainability",
    "financial_sustainability": "Financial sustainability",
    "team_capability":        "Team & capability",
    "market_context":         "Market context",
    "product_design":         "Product / service design",
    "supply_chain":           "Supply chain",
    "investor_alignment":     "Investor alignment",
}


def _category_label(key: str) -> str:
    if key in _CATEGORY_LABEL:
        return _CATEGORY_LABEL[key]
    # Sector-specific fallbacks — "sector_fintech" -> "Sector — Fintech"
    if key.startswith("sector_"):
        return "Sector — " + key.removeprefix("sector_").replace("_", " ").title()
    return key.replace("_", " ").title()


def _category_sort_key(cat: str) -> tuple[int, str]:
    """Rank categories by the natural DD sequence, sector_* always last."""
    if cat in _DD_SEQUENCE:
        return (_DD_SEQUENCE.index(cat), cat)
    if cat.startswith("sector_"):
        return (10_000, cat)
    return (1_000, cat)


def _question_sort_key(q: DDQuestion) -> tuple[int, int, str, str]:
    pri = _PRIORITY_RANK.get((q.priority or "medium").lower(), 1)
    cat_rank, cat = _category_sort_key(q.category)
    return (pri, cat_rank, cat, q.id or q.question)


def _evidence_asks_for(q: DDQuestion) -> list[str]:
    """Heuristic list of supporting docs to request for a given question."""
    text = (q.question + " " + (q.follow_up or "") + " " + " ".join(q.keywords or [])).lower()
    asks: list[str] = []
    for trigger, ask in _EVIDENCE_HINTS:
        if any(t.lower() in text for t in trigger):
            if ask not in asks:
                asks.append(ask)
    return asks


# ---------------------------------------------------------------------------
# Coverage helpers (shared with legacy renderer)
# ---------------------------------------------------------------------------

def _coverage_band(pct: float) -> tuple[str, str]:
    if pct >= 70:
        return "pass", "green"
    if pct >= 40:
        return "warn", "orange"
    return "fail", "red"


def _evidence_band(level: float) -> tuple[str, str]:
    if level >= 4:
        return "pass", "green"
    if level >= 2.5:
        return "warn", "orange"
    return "fail", "red"


def _category_rollup(
    addressed: Iterable[DDQuestionMatch],
    unanswered: Iterable[DDQuestion],
) -> list[tuple[str, int, int, float, int]]:
    """For each category, return (name, addressed, total, pct, high_gaps)."""
    addr: dict[str, int] = Counter()
    tot: dict[str, int] = Counter()
    high: dict[str, int] = Counter()
    for m in addressed:
        addr[m.question.category] += 1
        tot[m.question.category] += 1
    for q in unanswered:
        tot[q.category] += 1
        if (q.priority or "medium").lower() == "high":
            high[q.category] += 1
    rows: list[tuple[str, int, int, float, int]] = []
    for cat, total in tot.items():
        a = addr.get(cat, 0)
        pct = 100.0 * a / total if total else 0.0
        rows.append((cat, a, total, pct, high.get(cat, 0)))
    # Sort by severity: high-priority gaps first, then worst coverage
    rows.sort(key=lambda r: (-r[4], r[3], _category_sort_key(r[0])))
    return rows


def _priority_counts(items: Iterable[DDQuestion | DDQuestionMatch]) -> Counter:
    c: Counter = Counter()
    for it in items:
        q = it.question if isinstance(it, DDQuestionMatch) else it
        c[(q.priority or "medium").lower()] += 1
    return c


# ---------------------------------------------------------------------------
# HTML renderer — "DD Questionnaire Helper"
# ---------------------------------------------------------------------------

def render_dd_questionnaire_html(
    result: DDChecklistResult,
    *,
    company_name: str = "Company",
    document_label: str = "Source document",
    reviewer: str | None = None,
    run_at: str | None = None,
) -> str:
    """Render the questionnaire helper as a self-contained HTML page."""
    today = run_at or date.today().isoformat()

    cov_kind, cov_bar = _coverage_band(result.coverage_pct)
    ev_kind, _ = _evidence_band(result.avg_evidence_level)
    hp_count = len(result.high_priority_gaps)

    # --- Hero ------------------------------------------------------------
    hero = render_hero(
        eyebrow="Impact DD — Questionnaire Helper",
        title=escape(company_name),
        subtitle=(
            "Risk-first follow-up questionnaire generated from the Impact "
            f"Vision DD checklist, anchored to <b>{escape(document_label)}</b>. "
            "Send the <i>Information Request</i> section below to the "
            "founder / sponsor as the next DD step."
        ),
        meta=[
            ("Date",       today),
            ("Questions",  str(result.total_questions)),
            ("To ask",     str(len(result.unanswered))),
            ("Reviewer",   reviewer or "—"),
        ],
    )

    # --- KPIs ------------------------------------------------------------
    kpis = [
        {
            "label": "Checklist coverage",
            "value": f"{result.coverage_pct:.0f}%",
            "sub": f"{len(result.addressed)} of {result.total_questions} addressed",
            "kind": cov_kind,
            "badge": cov_kind.upper(), "badge_kind": cov_kind,
        },
        {
            "label": "High-priority gaps",
            "value": str(hp_count),
            "sub": "questions flagged priority=high",
            "kind": "fail" if hp_count >= 5 else "warn" if hp_count >= 1 else "pass",
            "badge": "ASK FIRST" if hp_count else "CLEAR",
            "badge_kind": "fail" if hp_count else "pass",
        },
        {
            "label": "Avg evidence level",
            "value": f"{result.avg_evidence_level:.1f}"
                     f"<span style='font-size:0.55em;color:var(--text-muted)'>/5</span>",
            "sub": "NESTA Standards of Evidence",
            "kind": ev_kind,
        },
        {
            "label": "Questions to send",
            "value": str(len(result.unanswered)),
            "sub": "full information request below",
            "kind": "neutral",
        },
    ]

    # --- Section: Key risk areas ---------------------------------------
    cat_rows = _category_rollup(result.addressed, result.unanswered)
    # Top 5 risky categories (≥1 high gap OR coverage < 40%)
    risk_rows = [r for r in cat_rows if r[4] > 0 or r[3] < 40.0][:5]
    if not risk_rows:
        risk_rows = cat_rows[:3]

    risk_cards: list[str] = []
    for cat, addr_n, total, pct_cat, hi in risk_rows:
        cov_colour = "green" if pct_cat >= 70 else "orange" if pct_cat >= 40 else "red"
        severity = "fail" if (hi >= 2 or pct_cat < 40) else "warn" if (hi >= 1 or pct_cat < 70) else "pass"
        sev_label = {"fail": "HIGH RISK", "warn": "MEDIUM", "pass": "LOW"}[severity]
        risk_cards.append(
            f'<div class="risk-card {severity}">'
            f'<div class="risk-head">'
            f'<div class="risk-cat">{escape(_category_label(cat))}</div>'
            f'<div class="risk-sev pill {severity}">{sev_label}</div>'
            f'</div>'
            f'<div class="risk-stats">'
            f'<span><b>{addr_n}/{total}</b> addressed</span> · '
            f'<span>{pct_cat:.0f}% coverage</span> · '
            + (f'<span style="color:var(--danger)"><b>{hi}</b> high-priority gap{"s" if hi != 1 else ""}</span>' if hi else '<span>no high-priority gaps</span>')
            + '</div>'
            f'<div class="bar-track" style="height:8px"><div class="bar-fill {cov_colour}" style="width:{int(max(0, min(100, pct_cat)))}%"></div></div>'
            f'</div>'
        )
    risk_html = (
        '<section class="card" id="risks">'
        '<h2 class="section-title">Key risk areas</h2>'
        '<p class="section-lede">Categories ranked by severity '
        '(count of high-priority gaps, then coverage percentage). '
        'These are the areas to interrogate first in the DD call.</p>'
        '<div class="risk-grid">' + "".join(risk_cards) + '</div>'
        '</section>'
    )

    # --- Section: Information Request (priority questionnaire) ---------
    # Sort unanswered by (priority, DD sequence, category, id)
    questions_sorted: list[DDQuestion] = sorted(result.unanswered, key=_question_sort_key)

    # Group by priority band for the callout
    by_priority: dict[str, list[DDQuestion]] = defaultdict(list)
    for q in questions_sorted:
        by_priority[(q.priority or "medium").lower()].append(q)

    q_sections: list[str] = []
    q_idx = 0
    for band in ("high", "medium", "low"):
        qs = by_priority.get(band, [])
        if not qs:
            continue
        band_class = {"high": "fail", "medium": "warn", "low": "na"}[band]
        band_label = {"high": "Priority 1 — Ask first", "medium": "Priority 2 — Ask in follow-up", "low": "Priority 3 — Nice to have"}[band]
        q_sections.append(
            f'<h3 class="q-band"><span class="pill {band_class}">'
            f'{len(qs)}</span> {band_label}</h3>'
        )
        # Group by category within the band so the analyst can skim
        by_cat: dict[str, list[DDQuestion]] = defaultdict(list)
        for q in qs:
            by_cat[q.category].append(q)
        for cat in sorted(by_cat, key=_category_sort_key):
            q_sections.append(
                f'<details class="q-cat" open>'
                f'<summary><span class="q-cat-label">{escape(_category_label(cat))}</span> '
                f'<span class="q-cat-count">{len(by_cat[cat])}</span></summary>'
            )
            for q in by_cat[cat]:
                q_idx += 1
                follow = (
                    f'<div class="q-followup"><b>Probe further:</b> {escape(q.follow_up)}</div>'
                    if q.follow_up else ''
                )
                asks = _evidence_asks_for(q)
                evidence_html = ''
                if asks:
                    evidence_html = (
                        '<div class="q-evidence"><b>Attach / ask for:</b><ul>'
                        + "".join(f'<li>{escape(a)}</li>' for a in asks)
                        + '</ul></div>'
                    )
                q_sections.append(
                    f'<div class="q-card">'
                    f'<div class="q-head">'
                    f'<span class="q-num">{q_idx}.</span>'
                    f'<span class="q-id">{escape(q.id or "")}</span>'
                    f'<span class="q-dim">{escape((q.dimension or "—").title())}</span>'
                    f'</div>'
                    f'<div class="q-text">{escape(q.question)}</div>'
                    f'{follow}'
                    f'{evidence_html}'
                    f'<div class="q-answer"><b>Founder response:</b><div class="q-answer-slot">&nbsp;</div></div>'
                    f'</div>'
                )
            q_sections.append('</details>')

    if not questions_sorted:
        q_sections = [
            '<div class="callout ok">Every checklist question is already '
            'addressed by the supplied document — no information request '
            'needed at this stage.</div>'
        ]

    info_request_html = (
        '<section class="card" id="questionnaire">'
        '<h2 class="section-title">Information request (questionnaire to send)</h2>'
        '<p class="section-lede">Ordered first by severity, then by the '
        'natural DD sequence (Thesis → ToC → What → Who → How-much → '
        'Contribution → Risk → Measurement → Governance → Sector → Exit). '
        'Each question lists the supporting documents to request alongside '
        'the answer.</p>'
        + "".join(q_sections)
        + '<p class="questionnaire-footer">Tip: use <code>'
        'ImpactVision().render_dd_questionnaire_docx(...)</code> '
        'to export an editable Word (.docx) version of this section.</p>'
        '</section>'
    )

    # --- Section: Evidence / document gaps -----------------------------
    doc_counter: Counter = Counter()
    doc_examples: dict[str, list[str]] = defaultdict(list)
    for q in questions_sorted:
        for a in _evidence_asks_for(q):
            doc_counter[a] += 1
            if len(doc_examples[a]) < 2:
                doc_examples[a].append(q.question)

    if doc_counter:
        rows_html = "".join(
            f'<tr>'
            f'<td><b>{escape(doc)}</b><div style="font-size:0.82em;color:var(--text-secondary);margin-top:2px">'
            f'e.g. &ldquo;{escape(doc_examples[doc][0])}&rdquo;</div></td>'
            f'<td><span class="pill warn">{count}</span></td>'
            f'<td><label><input type="checkbox"> Requested</label><br>'
            f'<label><input type="checkbox"> Received</label></td>'
            f'</tr>'
            for doc, count in doc_counter.most_common()
        )
        evidence_html = (
            '<section class="card" id="evidence-gaps">'
            '<h2 class="section-title">Evidence &amp; document gaps</h2>'
            '<p class="section-lede">Consolidated list of supporting documents '
            'the founder should provide. Each row aggregates all questions that '
            'need the same class of evidence, so you can request them in a '
            'single email.</p>'
            '<table class="data"><thead><tr>'
            '<th>Document / evidence</th>'
            '<th># questions</th>'
            '<th>Status</th>'
            '</tr></thead><tbody>' + rows_html + '</tbody></table>'
            '</section>'
        )
    else:
        evidence_html = (
            '<section class="card" id="evidence-gaps">'
            '<h2 class="section-title">Evidence &amp; document gaps</h2>'
            '<div class="callout ok">No additional supporting documents required '
            '— all unanswered questions can be answered directly by the founder.</div>'
            '</section>'
        )

    # --- Section: Overview / coverage snapshot -------------------------
    pct = max(0, min(100, int(result.coverage_pct)))
    overview = (
        '<section class="card" id="overview">'
        '<h2 class="section-title">Coverage snapshot</h2>'
        '<p class="section-lede">Where the submitted document already '
        'answers the checklist. Full category breakdown is further below.</p>'
        '<h3>Overall coverage</h3>'
        '<div class="bar-track" style="height:14px">'
        f'<div class="bar-fill {cov_bar}" style="width:{pct}%"></div></div>'
        f'<p style="margin-top:6px;font-size:0.88em;color:var(--text-secondary)">'
        f'<b>{result.coverage_pct:.1f}%</b> of {result.total_questions} '
        f'checklist questions addressed · avg evidence level '
        f'<b>{result.avg_evidence_level:.1f}/5</b>.</p>'
        '</section>'
    )

    # --- Section: Category roll-up -------------------------------------
    cat_body = ""
    for cat, addr_n, total, pct_cat, hi in cat_rows:
        p = max(0, min(100, int(pct_cat)))
        cov_colour = "green" if pct_cat >= 70 else "orange" if pct_cat >= 40 else "red"
        hi_cell = (
            f'<span class="pill fail">{hi}</span>' if hi
            else '<span style="color:var(--text-secondary)">—</span>'
        )
        cat_body += (
            f'<tr><td><b>{escape(_category_label(cat))}</b></td>'
            f'<td>{addr_n}/{total}</td>'
            f'<td style="min-width:180px"><div class="bar-track">'
            f'<div class="bar-fill {cov_colour}" style="width:{p}%"></div></div></td>'
            f'<td>{pct_cat:.0f}%</td>'
            f'<td>{hi_cell}</td></tr>'
        )
    category_html = (
        '<section class="card" id="categories">'
        '<h2 class="section-title">Coverage by category</h2>'
        '<p class="section-lede">Worst-covered categories with high-priority '
        'gaps float to the top.</p>'
        '<table class="data"><thead><tr>'
        '<th>Category</th><th>Addressed / Total</th><th>Coverage</th><th>%</th>'
        '<th>High-priority gaps</th>'
        '</tr></thead><tbody>' + cat_body + '</tbody></table></section>'
    )

    # --- Section: Addressed questions (appendix) -----------------------
    by_level: dict[int, list[DDQuestionMatch]] = defaultdict(list)
    for m in result.addressed:
        by_level[m.evidence_level].append(m)

    addressed_inner: list[str] = []
    if not result.addressed:
        addressed_inner.append(
            '<div class="callout warn">No checklist questions appear to be '
            'addressed in the supplied document.</div>'
        )
    else:
        for lvl in sorted(by_level.keys(), reverse=True):
            matches = by_level[lvl]
            band_kind, _ = _evidence_band(lvl)
            addressed_inner.append(
                f'<h3>Level {lvl} '
                f'<span class="pill {band_kind}">{len(matches)}</span> '
                f'<span style="font-weight:400;color:var(--text-secondary);font-size:0.9em">'
                f'· {escape(EVIDENCE_LEVELS.get(lvl, ""))}</span></h3>'
            )
            rows = "".join(
                f'<tr><td style="max-width:520px">{escape(m.question.question)}</td>'
                f'<td style="font-size:0.85em;color:var(--text-secondary)">{escape(_category_label(m.question.category))}</td>'
                f'<td>{int(m.confidence * 100)}%</td>'
                f'<td style="font-size:0.85em;color:var(--text-secondary)">'
                f'{", ".join(escape(k) for k in m.matched_keywords[:4])}</td></tr>'
                for m in matches
            )
            addressed_inner.append(
                '<table class="data"><thead><tr>'
                '<th>Question</th><th>Category</th><th>Confidence</th><th>Keywords</th>'
                '</tr></thead><tbody>' + rows + '</tbody></table>'
            )
    addressed_html = (
        '<section class="card" id="addressed">'
        '<details>'
        '<summary><h2 class="section-title" style="display:inline-block;margin:0">'
        'Questions already addressed (appendix)</h2></summary>'
        '<p class="section-lede">Audit-trail of what the auto-matcher picked '
        'up, grouped by NESTA evidence level. Inspect if you want to sanity-'
        'check any match before approving.</p>'
        + "".join(addressed_inner) +
        '</details>'
        '</section>'
    )

    # --- Section: Evidence legend --------------------------------------
    legend_rows = "".join(
        f'<tr><td><b>{lvl}</b></td><td>{escape(label)}</td></tr>'
        for lvl, label in EVIDENCE_LEVELS.items()
    )
    legend_html = (
        '<section class="card" id="legend">'
        '<h2 class="section-title">Evidence level reference</h2>'
        '<p class="section-lede">NESTA Standards of Evidence — used to grade '
        'every matched question.</p>'
        '<table class="data"><thead><tr><th>Level</th><th>Description</th></tr></thead><tbody>'
        + legend_rows + '</tbody></table></section>'
    )

    toc = render_toc([
        ("risks",          "1. Key risk areas"),
        ("questionnaire",  "2. Information request"),
        ("evidence-gaps",  "3. Evidence & document gaps"),
        ("overview",       "4. Coverage snapshot"),
        ("categories",     "5. By category"),
        ("addressed",      "6. Addressed (appendix)"),
        ("legend",         "7. Evidence levels"),
    ])

    extra_css = """
    <style>
      .risk-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 14px; margin-top: 10px; }
      .risk-card { padding: 14px 16px; border-radius: 10px; border: 1px solid var(--border); background: var(--surface); box-shadow: var(--shadow-sm); border-left-width: 4px; }
      .risk-card.fail { border-left-color: var(--danger); }
      .risk-card.warn { border-left-color: var(--warning); }
      .risk-card.pass { border-left-color: var(--success); }
      .risk-head { display: flex; justify-content: space-between; align-items: center; gap: 10px; }
      .risk-cat { font-weight: 700; color: var(--text); }
      .risk-sev { font-size: 0.72em; font-weight: 700; letter-spacing: 0.04em; }
      .risk-stats { font-size: 0.85em; color: var(--text-secondary); margin: 6px 0 8px; line-height: 1.45; }

      .q-band { margin-top: 18px; padding-bottom: 6px; border-bottom: 1px solid var(--border); }
      .q-cat { margin: 10px 0; border: 1px solid var(--border); border-radius: 8px; background: var(--bg); }
      .q-cat > summary { padding: 10px 14px; cursor: pointer; font-weight: 600; color: var(--text); display: flex; justify-content: space-between; align-items: center; list-style: none; }
      .q-cat > summary::-webkit-details-marker { display: none; }
      .q-cat > summary::before { content: "▸"; margin-right: 8px; color: var(--text-secondary); transition: transform 0.15s; }
      .q-cat[open] > summary::before { transform: rotate(90deg); display: inline-block; }
      .q-cat-label { flex: 1; }
      .q-cat-count { font-size: 0.75em; background: var(--primary-light); color: var(--primary); padding: 2px 8px; border-radius: 10px; font-weight: 700; }

      .q-card { margin: 10px 14px 14px; padding: 12px 14px; background: var(--surface); border: 1px solid var(--border); border-radius: 8px; }
      .q-head { display: flex; gap: 8px; align-items: center; margin-bottom: 6px; font-size: 0.78em; color: var(--text-secondary); text-transform: uppercase; letter-spacing: 0.04em; }
      .q-num { font-weight: 800; color: var(--primary); }
      .q-id { background: var(--bg); padding: 1px 6px; border-radius: 4px; font-family: monospace; }
      .q-dim { margin-left: auto; background: var(--primary-light); color: var(--primary); padding: 1px 7px; border-radius: 10px; font-weight: 600; }
      .q-text { font-weight: 600; color: var(--text); line-height: 1.5; }
      .q-followup { margin-top: 6px; padding: 6px 10px; background: var(--bg); border-left: 3px solid var(--warning); border-radius: 0 4px 4px 0; font-size: 0.88em; color: var(--text-secondary); }
      .q-evidence { margin-top: 8px; font-size: 0.85em; color: var(--text); }
      .q-evidence ul { margin: 4px 0 0 18px; padding: 0; }
      .q-evidence li { margin: 2px 0; color: var(--text-secondary); }
      .q-answer { margin-top: 10px; font-size: 0.85em; color: var(--text-secondary); }
      .q-answer-slot { margin-top: 4px; min-height: 44px; border: 1px dashed var(--border); border-radius: 6px; padding: 8px; background: #fafbfc; }

      .questionnaire-footer { margin-top: 14px; font-size: 0.85em; color: var(--text-muted); font-style: italic; }
      .questionnaire-footer code { background: var(--bg); padding: 1px 6px; border-radius: 4px; font-size: 0.92em; }

      @media print {
        .q-cat > summary::before { display: none; }
        details { display: block !important; }
        summary { pointer-events: none; }
        .q-card { break-inside: avoid; }
        .risk-card { break-inside: avoid; }
      }
    </style>
    """

    body = (
        extra_css +
        '<div class="page">'
        f'{toc}'
        '<main>'
        f'{hero}'
        f'{render_kpi_strip(kpis)}'
        f'{risk_html}'
        f'{info_request_html}'
        f'{evidence_html}'
        f'{overview}'
        f'{category_html}'
        f'{addressed_html}'
        f'{legend_html}'
        f'{render_footer("DD Questionnaire Helper — review with the investment team before sending.")}'
        '</main>'
        '</div>'
    )

    return wrap_document(
        title=f"DD Questionnaire Helper — {company_name}",
        body_html=body,
    )


def save_dd_questionnaire_html(
    result: DDChecklistResult,
    path: str | Path,
    **kwargs,
) -> Path:
    """Render and persist the DD questionnaire helper HTML, returning its Path."""
    html = render_dd_questionnaire_html(result, **kwargs)
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(html, encoding="utf-8")
    return p


# ---------------------------------------------------------------------------
# Word (.docx) renderer
# ---------------------------------------------------------------------------

def render_dd_questionnaire_docx(
    result: DDChecklistResult,
    path: str | Path,
    *,
    company_name: str = "Company",
    document_label: str = "Source document",
    reviewer: str | None = None,
    run_at: str | None = None,
) -> Path:
    """Render the DD Questionnaire Helper as an editable ``.docx`` file.

    Requires the optional dependency ``python-docx``. Raises ``ImportError``
    with a clear install hint otherwise.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ImportError as exc:  # pragma: no cover — optional
        raise ImportError(
            "Word export requires the optional dependency 'python-docx'. "
            "Install with: pip install python-docx"
        ) from exc

    today = run_at or date.today().isoformat()
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # --- Cover block ----------------------------------------------------
    doc.add_heading("Impact DD — Questionnaire Helper", level=0)
    doc.add_paragraph(f"Company: {company_name}")
    doc.add_paragraph(f"Source document: {document_label}")
    doc.add_paragraph(f"Prepared: {today}   ·   Reviewer: {reviewer or '—'}")

    # --- Executive summary ---------------------------------------------
    doc.add_heading("Executive summary", level=1)
    summary = doc.add_paragraph()
    summary.add_run(
        f"Checklist coverage: {result.coverage_pct:.1f}% "
        f"({len(result.addressed)} of {result.total_questions} addressed). "
    )
    summary.add_run(
        f"High-priority gaps to raise first: {len(result.high_priority_gaps)}. "
    )
    summary.add_run(
        f"Average NESTA evidence level of addressed questions: "
        f"{result.avg_evidence_level:.1f}/5. "
    )
    summary.add_run(
        "The following questionnaire is ordered by severity, then by the "
        "natural DD sequence. Fill the 'Founder response' boxes as answers "
        "come back, and tick off the supporting documents as you receive them."
    )

    # --- Key risk areas -------------------------------------------------
    cat_rows = _category_rollup(result.addressed, result.unanswered)
    risk_rows = [r for r in cat_rows if r[4] > 0 or r[3] < 40.0][:5] or cat_rows[:3]
    doc.add_heading("Key risk areas", level=1)
    for cat, addr_n, total, pct_cat, hi in risk_rows:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(_category_label(cat))
        run.bold = True
        para.add_run(
            f" — {addr_n}/{total} addressed · {pct_cat:.0f}% coverage"
            + (f" · {hi} high-priority gap" + ("s" if hi != 1 else "") if hi else "")
        )

    # --- Information request -------------------------------------------
    doc.add_heading("Information request (questionnaire to send)", level=1)
    doc.add_paragraph(
        "Sorted by priority and DD sequence. Each block has space for "
        "the founder's response and a checklist of supporting documents."
    )

    questions_sorted = sorted(result.unanswered, key=_question_sort_key)
    by_priority: dict[str, list[DDQuestion]] = defaultdict(list)
    for q in questions_sorted:
        by_priority[(q.priority or "medium").lower()].append(q)

    q_idx = 0
    for band in ("high", "medium", "low"):
        qs = by_priority.get(band, [])
        if not qs:
            continue
        band_label = {
            "high":   "Priority 1 — Ask first",
            "medium": "Priority 2 — Ask in follow-up",
            "low":    "Priority 3 — Nice to have",
        }[band]
        doc.add_heading(band_label + f" ({len(qs)} questions)", level=2)

        by_cat: dict[str, list[DDQuestion]] = defaultdict(list)
        for q in qs:
            by_cat[q.category].append(q)
        for cat in sorted(by_cat, key=_category_sort_key):
            doc.add_heading(_category_label(cat), level=3)
            for q in by_cat[cat]:
                q_idx += 1
                # Question header
                qp = doc.add_paragraph()
                num_run = qp.add_run(f"{q_idx}. ")
                num_run.bold = True
                num_run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
                id_run = qp.add_run(f"[{q.id}] ")
                id_run.font.size = Pt(9)
                id_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                main_run = qp.add_run(q.question)
                main_run.bold = True

                if q.follow_up:
                    fp = doc.add_paragraph()
                    fr = fp.add_run("Probe further: ")
                    fr.italic = True
                    fr.bold = True
                    fp.add_run(q.follow_up).italic = True

                asks = _evidence_asks_for(q)
                if asks:
                    doc.add_paragraph("Attach / ask for:", style="Intense Quote")
                    for ask in asks:
                        doc.add_paragraph(ask, style="List Bullet")

                resp = doc.add_paragraph()
                rr = resp.add_run("Founder response:")
                rr.bold = True
                resp_box = doc.add_paragraph(" ")
                resp_box.paragraph_format.left_indent = Pt(18)
                resp_box.paragraph_format.space_after = Pt(12)

    if not questions_sorted:
        doc.add_paragraph(
            "No outstanding questions — every checklist item is already "
            "addressed by the supplied document."
        )

    # --- Evidence / document gaps --------------------------------------
    doc.add_heading("Evidence & document gaps (consolidated)", level=1)
    doc_counter: Counter = Counter()
    for q in questions_sorted:
        for a in _evidence_asks_for(q):
            doc_counter[a] += 1
    if doc_counter:
        for ask, count in doc_counter.most_common():
            bullet = doc.add_paragraph(style="List Bullet")
            r = bullet.add_run(f"☐ {ask}")
            r.bold = True
            bullet.add_run(f"  ({count} related question" + ("s" if count != 1 else "") + ")")
    else:
        doc.add_paragraph("No additional supporting documents required.")

    # --- Sign-off block ------------------------------------------------
    doc.add_heading("Sign-off", level=1)
    doc.add_paragraph(f"Prepared by: {reviewer or '_____________________'}")
    doc.add_paragraph("Reviewed by: _____________________")
    doc.add_paragraph("Date reviewed: _____________________")

    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# Backwards-compatible aliases (legacy DD coverage report API)
# ---------------------------------------------------------------------------

#: Historical name — kept so downstream callers keep working. Renders the
#: new DD Questionnaire Helper HTML.
render_dd_report_html = render_dd_questionnaire_html
save_dd_report_html = save_dd_questionnaire_html


__all__ = [
    "render_dd_questionnaire_html",
    "save_dd_questionnaire_html",
    "render_dd_questionnaire_docx",
    # Legacy aliases
    "render_dd_report_html",
    "save_dd_report_html",
]
